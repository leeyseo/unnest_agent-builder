"""DOCX 파서 (python-docx) — 문단 + 표를 추출한다."""
from __future__ import annotations

from agentsdk import Block, Component, NormalizedDocument, RawFile, port
from agentengine import BAD_INPUT, EngineError


class DOCXParser(Component):
    """워드(.docx) 문서에서 문단·표 블록을 추출한다."""

    display_name = "DOCX 파서"
    category = "parsers"
    icon = "file-text"

    file: RawFile = port(input=True, display_name="DOCX 파일")
    document: NormalizedDocument = port(output=True, display_name="문서")

    def run(self) -> NormalizedDocument:
        import docx

        if self.file is None:
            raise EngineError("파일 입력이 연결되지 않았습니다.", BAD_INPUT)
        try:
            doc = docx.Document(self.file.path)
        except Exception as ex:
            raise EngineError(
                f"DOCX를 열 수 없습니다: {self.file.filename} — {ex}. "
                ".doc(구형식)은 지원하지 않습니다. .docx로 저장해 주세요.",
                BAD_INPUT,
            ) from ex

        blocks: list[Block] = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                blocks.append(Block(type="text", content=text, meta={"para": i + 1}))
        for t, table in enumerate(doc.tables):
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells)
                for row in table.rows
            ]
            content = "\n".join(r for r in rows if r.strip(" |"))
            if content:
                blocks.append(Block(type="table", content=content, meta={"table": t + 1}))

        if not blocks:
            raise EngineError(f"'{self.file.filename}'에서 내용을 찾지 못했습니다.", BAD_INPUT)
        return NormalizedDocument(doc_type="docx", source=self.file.filename, blocks=blocks)
